"""
DOCX file processor for extracting text from Word documents.
"""

import logging
import time
from pathlib import Path
from docx import Document

from ..core.base import BaseProcessor
from ..core.types import ProcessorResult


# Setup logging
logger = logging.getLogger(__name__)


class DOCXProcessor(BaseProcessor):
    """Extract text from DOCX files using python-docx."""

    def can_process(self, file_path: Path) -> bool:
        """Check if this processor can handle the file."""
        return file_path.suffix.lower() in ['.docx', '.doc']

    def process(self, file_path: Path, output_dir: Path) -> ProcessorResult:
        """Process DOCX file and extract text."""
        start_time = time.time()
        output_file = output_dir / f"{file_path.stem}.txt"

        try:
            # Skip if output already exists
            if output_file.exists():
                logger.info(f"Skipping (already exists): {file_path.name}")
                return ProcessorResult(
                    success=True,
                    text="",
                    source_file=file_path,
                    output_file=output_file,
                    processor_type="docx",
                    processing_time=0,
                    error_message="File already exists"
                )

            # Handle .doc files (older format)
            if file_path.suffix.lower() == '.doc':
                logger.warning(f"Old .doc format detected for {file_path.name}. Consider converting to .docx for better results.")

            # Extract text from DOCX
            doc = Document(file_path)

            # Extract structured text preserving formatting
            text = self._extract_structured_text(doc)

            # Check if we extracted any meaningful text
            if not text or len(text.strip()) < 5:
                logger.warning(f"Very little text extracted from {file_path.name}")
                text = "[WARNING: Very little text found in document]"

            # Save to output file
            output_file.parent.mkdir(parents=True, exist_ok=True)
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)

            duration = time.time() - start_time
            logger.info(f"✅ Completed: {file_path.name} ({duration:.1f}s, {len(text)} chars)")

            return ProcessorResult(
                success=True,
                text=text,
                source_file=file_path,
                output_file=output_file,
                processor_type="docx",
                processing_time=duration
            )

        except Exception as e:
            duration = time.time() - start_time
            error_msg = f"Failed to process DOCX {file_path.name}: {e}"
            logger.error(error_msg)

            # Special handling for .doc files
            if file_path.suffix.lower() == '.doc':
                error_msg += " (Note: .doc files may require conversion to .docx)"

            return ProcessorResult(
                success=False,
                text="",
                source_file=file_path,
                output_file=None,
                processor_type="docx",
                processing_time=duration,
                error_message=str(e)
            )

    def _extract_structured_text(self, doc):
        """Extract text while preserving document structure."""
        structured_lines = []

        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                structured_lines.append("")  # Preserve empty lines
                continue

            formatted_line = self._process_paragraph_with_structure(paragraph)
            if formatted_line:
                structured_lines.append(formatted_line)

        # Add tables at the end if they exist
        table_text = self._extract_table_text(doc)
        if table_text:
            structured_lines.extend(["", "--- TABLES ---", ""] + table_text)

        return "\n".join(structured_lines)

    def _process_paragraph_with_structure(self, paragraph):
        """Process paragraph preserving bullets, numbers, and headings."""
        text = paragraph.text.strip()
        if not text:
            return None

        # Check for heading styles
        if paragraph.style.name.startswith('Heading'):
            try:
                level = int(paragraph.style.name.split()[-1])
                return f"{'#' * level} {text}"
            except (ValueError, IndexError):
                return f"# {text}"  # Default to H1 if parsing fails

        # Check for Word's built-in list formatting
        if paragraph._element.xpath('.//w:numPr'):
            # Get indentation level for nested lists
            level = self._get_list_level(paragraph)
            indent = "  " * level

            # Check if it's numbered or bulleted
            if self._is_numbered_list(paragraph):
                return f"{indent}1. {text}"  # Use generic numbering
            else:
                return f"{indent}• {text}"   # Use bullet

        # Regular paragraph
        return text

    def _get_list_level(self, paragraph):
        """Get the indentation level of a list item."""
        try:
            num_element = paragraph._element.xpath('.//w:numPr')[0]
            level_elem = num_element.xpath('.//w:ilvl')
            if level_elem:
                return int(level_elem[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))
        except (IndexError, ValueError, AttributeError):
            pass
        return 0

    def _is_numbered_list(self, paragraph):
        """Simple check if this should be numbered vs bulleted."""
        # This is a basic implementation - Word's numbering detection is complex
        # For now, assume most lists are bulleted unless we detect specific numbering
        try:
            num_element = paragraph._element.xpath('.//w:numPr')[0]
            num_id_elem = num_element.xpath('.//w:numId')
            if num_id_elem:
                # Could check numbering definition, but for simplicity use bullets
                return False
        except (IndexError, AttributeError):
            pass
        return False

    def _extract_table_text(self, doc):
        """Extract and format table text."""
        table_lines = []
        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    table_lines.append(" | ".join(row_cells))
        return table_lines